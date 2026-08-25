"""Build thesis_master.docx from markdown (part 2: main parser)."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches

from md2docx_lib import PERSIAN, add_heading, add_image, add_para, add_table

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "thesis" / "fa" / "thesis_master.md"
IMGDIR = SRC.parent


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.9)

    skip = False
    table_buf, math_buf = [], []
    i = 0

    def flush_table():
        nonlocal table_buf
        if len(table_buf) >= 2:
            add_table(doc, table_buf)
        table_buf = []

    while i < len(lines):
        ln = lines[i]
        s = ln.strip()

        if s.startswith("## وضعیت نگارش"):
            skip = True
        if skip and s.startswith("## فهرست شکل‌ها"):
            skip = False
        if skip:
            i += 1
            continue

        if s.startswith("|"):
            table_buf.append(s)
            i += 1
            continue
        flush_table()

        if not s or s == "---":
            i += 1
            continue

        m = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", s)
        if m:
            add_image(doc, m.group(1), m.group(2), IMGDIR)
            i += 1
            continue

        m = re.match(r"^\$\$(.+)\$\$$", s)
        if m:
            add_para(doc, m.group(1).strip(), align_right=False, size=11)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            add_heading(doc, m.group(2).strip(), min(len(m.group(1)), 3))
            i += 1
            continue

        if s.startswith(">"):
            content = s.lstrip(">").strip()
            if content:
                per = bool(PERSIAN.search(content))
                p = add_para(doc, "", align_right=per, size=13)
                from md2docx_lib import add_inline
                add_inline(p, content, base_persian=per, size=13)
            i += 1
            continue

        m = re.match(r"^[-*]\s+(.*)$", s)
        if m:
            add_para(doc, "• " + m.group(1), size=12)
            i += 1
            continue

        m = re.match(r"^(\d+)[.-]\s+(.*)$", s)
        if m:
            add_para(doc, f"{m.group(1)}. " + m.group(2),
                     align_right=bool(PERSIAN.search(m.group(2))), size=12)
            i += 1
            continue

        per = bool(PERSIAN.search(s))
        add_para(doc, s, align_right=per, size=12.5 if len(s) < 60 else 12)
        i += 1

    flush_table()
    out = SRC.with_suffix(".docx")
    doc.save(out)
    print("saved ->", out)


if __name__ == "__main__":
    main()
