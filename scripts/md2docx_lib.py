"""RTL-aware markdown helpers for DOCX/PPTX builders (part 1: docx helpers)."""
import re
from pathlib import Path

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PERSIAN = re.compile(r"[\u0600-\u06FF]")


def set_rtl_para(p):
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(pPr.makeelement(qn("w:bidi"), {}))


def style_run(run, persian=False, bold=False, italic=False, code=False, size=None):
    run.bold = bold or None
    run.italic = italic or None
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {}); rpr.append(rf)
    rf.set(qn("w:ascii"), "Consolas" if code else "Times New Roman")
    rf.set(qn("w:hAnsi"), "Consolas" if code else "Times New Roman")
    rf.set(qn("w:cs"), "Consolas" if code else "B Nazanin")
    run.font.size = Pt(size or (10 if code else 12))
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = rpr.makeelement(qn("w:szCs"), {}); rpr.append(szcs)
    szcs.set(qn("w:val"), str(int((size or (10 if code else 13)) * 2)))
    if persian and rpr.find(qn("w:rtl")) is None:
        rpr.append(rpr.makeelement(qn("w:rtl"), {}))


TOKEN = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+`)")


def add_inline(p, text, base_persian=True, size=None):
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); style_run(r, base_persian, bold=True, size=size)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = p.add_run(part[1:-1]); style_run(r, code=True,
                                                 size=(size - 1.5) if size else None)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); style_run(r, base_persian, italic=True, size=size)
        else:
            part = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", part)
            per = base_persian or bool(PERSIAN.search(part))
            r = p.add_run(part); style_run(r, per, size=size)


def add_para(doc, text="", align_right=True, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT
    if align_right:
        set_rtl_para(p)
    if text:
        add_inline(p, text, base_persian=align_right, size=size)
    return p


def add_heading(doc, text, level):
    sizes = {1: 18, 2: 15, 3: 13.5}
    p = add_para(doc, align_right=True, size=sizes.get(level, 12))
    add_inline(p, text, size=sizes.get(level, 12))
    for r in p.runs:
        r.bold = True
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)
        rpr = r._element.get_or_add_rPr()
        if rpr.find(qn("w:bCs")) is None:
            rpr.append(rpr.makeelement(qn("w:bCs"), {}))
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in r.strip("|").split("|")] for r in rows[2:]]
    t = doc.add_table(rows=len(body) + 1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblpr = t._tbl.tblPr
    if tblpr.find(qn("w:bidiVisual")) is None:
        tblpr.append(tblpr.makeelement(qn("w:bidiVisual"), {}))

    def fill(cell, text, bold=False, right=False):
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right else WD_ALIGN_PARAGRAPH.CENTER
        add_inline(cp, text, size=10.5)
        if bold:
            for r in cp.runs:
                r.bold = True

    for j, ctext in enumerate(header):
        fill(t.rows[0].cells[j], ctext, bold=True)
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            val = row[j] if j < len(row) else ""
            fill(t.rows[i].cells[j], val, right=(j == 0))
    doc.add_paragraph()


def add_image(doc, alt, rel_path, basedir, width=5.8):
    img = Path(basedir) / rel_path
    if not img.exists():
        add_para(doc, f"[image missing: {rel_path}]", align_right=False)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Inches(width))
    cap = add_para(doc, alt, align_right=bool(PERSIAN.search(alt)), size=10.5)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
